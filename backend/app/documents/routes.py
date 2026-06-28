import os
import uuid
import csv
from datetime import datetime
from typing import List, Optional
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, BackgroundTasks, status
from pydantic import BaseModel
from sqlalchemy.orm import Session
from app.database.session import get_db, SessionLocal
from app.documents.models import Document
from app.auth.models import User
from app.auth.security import get_current_user
from app.products.routes import verify_business_access

# Text extraction library imports (imported inside extraction helper to prevent import-time failure if packages are missing)
# we will import:
# from pypdf import PdfReader
# from docx import Document as DocxDocument

router = APIRouter()

UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "uploads")
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
ALLOWED_EXTENSIONS = {".txt", ".csv", ".pdf", ".docx"}

# --- Pydantic Schemas ---

class DocumentResponse(BaseModel):
    id: uuid.UUID
    business_id: uuid.UUID
    file_name: str
    file_type: str
    file_path: str
    processed_status: str
    extracted_text: Optional[str] = None
    created_at: datetime
    updated_at: datetime

    class Config:
        from_attributes = True


# --- Extraction Helpers ---

def extract_text_from_file(file_path: str, file_type: str) -> str:
    ext = file_type.lower()
    if ext == '.txt':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
            
    elif ext == '.csv':
        output = []
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f)
            for idx, row in enumerate(reader):
                # Clean elements and skip empty rows
                clean_row = [cell.strip() for cell in row if cell.strip()]
                if clean_row:
                    output.append(f"Row {idx + 1}: " + ", ".join(clean_row))
        return "\n".join(output)
        
    elif ext == '.pdf':
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text_content = []
        for idx, page in enumerate(reader.pages):
            t = page.extract_text()
            if t and t.strip():
                text_content.append(t.strip())
        return "\n\n".join(text_content)
        
    elif ext == '.docx':
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        # Parse tables in docx
        for t_idx, table in enumerate(doc.tables):
            table_lines = []
            for r_idx, row in enumerate(table.rows):
                row_text = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                # De-duplicate adjacent identical cells due to cell merging in docx
                dedup_row = []
                for cell in row_text:
                    if not dedup_row or dedup_row[-1] != cell:
                        dedup_row.append(cell)
                if dedup_row:
                    table_lines.append(f"Row {r_idx + 1}: " + " | ".join(dedup_row))
            if table_lines:
                text_content.append(f"--- Table {t_idx + 1} ---\n" + "\n".join(table_lines))
        return "\n\n".join(text_content)
        
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def process_document_background(document_id: uuid.UUID):
    """Background task to extract text and update the document status in the DB."""
    db = SessionLocal()
    try:
        doc = db.query(Document).filter(Document.id == document_id).first()
        if not doc:
            return
        
        extracted = extract_text_from_file(doc.file_path, doc.file_type)
        
        doc.extracted_text = extracted
        doc.processed_status = "processed"
        db.commit()
    except Exception as e:
        print(f"Error extracting text from document {document_id}: {e}")
        db.rollback()
        # Refresh and update status to failed
        doc = db.query(Document).filter(Document.id == document_id).first()
        if doc:
            doc.processed_status = "failed"
            db.commit()
    finally:
        db.close()


# --- API Routes ---

@router.post("/businesses/{business_id}/documents/upload", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def upload_document(
    business_id: uuid.UUID,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Uploads a document under a business profile. Scoped to authorized owner/staff/admin."""
    verify_business_access(business_id, db, current_user)
    
    # 1. Validate File Extension
    file_name = file.filename
    _, ext = os.path.splitext(file_name)
    ext = ext.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file format. Supported formats: {', '.join(ALLOWED_EXTENSIONS)}"
        )
        
    # 2. Validate File Size
    # Read the file contents to measure size
    contents = await file.read()
    file_size = len(contents)
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File size exceeds the 5MB limit."
        )
        
    # Reset seek pointer for any subsequent file reading
    await file.seek(0)
    
    # 3. Create Business Upload Directory
    business_dir = os.path.join(UPLOAD_DIR, str(business_id))
    os.makedirs(business_dir, exist_ok=True)
    
    # 4. Generate Unique Safe File Name
    unique_filename = f"{uuid.uuid4()}{ext}"
    file_path = os.path.join(business_dir, unique_filename)
    
    # 5. Save file to disk
    try:
        with open(file_path, "wb") as f:
            f.write(contents)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to save file to disk: {str(e)}"
        )
        
    # 6. Save document record in database
    new_doc = Document(
        business_id=business_id,
        file_name=file_name,
        file_type=ext,
        file_path=file_path,
        processed_status="pending"
    )
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)
    
    # 7. Spawn Background Processing Task for text extraction
    background_tasks.add_task(process_document_background, new_doc.id)
    
    return new_doc


@router.get("/businesses/{business_id}/documents", response_model=List[DocumentResponse])
def list_documents(
    business_id: uuid.UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Lists all uploaded documents for a business. Scoped to authorized owner/staff/admin."""
    verify_business_access(business_id, db, current_user)
    return db.query(Document).filter(Document.business_id == business_id).order_by(Document.created_at.desc()).all()


@router.get("/documents/{document_id}", response_model=DocumentResponse)
def get_document(
    document_id: uuid.UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Retrieves metadata and status of a single document. Scoped to authorized owner/staff/admin."""
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
        
    verify_business_access(doc.business_id, db, current_user)
    return doc


@router.delete("/documents/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_document(
    document_id: uuid.UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Deletes a document record and removes its physical file. Scoped to authorized owner/staff/admin."""
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
        
    verify_business_access(doc.business_id, db, current_user)
    
    # Delete the physical file if it exists
    if os.path.exists(doc.file_path):
        try:
            os.remove(doc.file_path)
        except Exception as e:
            print(f"Warning: Failed to delete physical file {doc.file_path}: {e}")
            
    db.delete(doc)
    db.commit()
    return None
